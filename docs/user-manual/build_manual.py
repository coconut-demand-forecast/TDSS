# -*- coding: utf-8 -*-
"""Builds the TDSS user manual (Word .docx) from structured content + screenshots.
Run with: venv/Scripts/python.exe docs/user-manual/build_manual.py
Screenshots are looked up in docs/user-manual/images/<key>.png; missing ones
get a placeholder note instead of failing the build.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, "images")
OUT_PATH = os.path.join(BASE_DIR, "TDSS_คู่มือการใช้งาน.docx")

ACCENT = RGBColor(0xD7, 0x19, 0x20)
DARK = RGBColor(0x17, 0x17, 0x17)

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Tahoma"
style.font.size = Pt(11)


def add_title_page():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\nTDSS")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("ระบบวางแผนงานขนส่ง")
    run2.font.size = Pt(22)
    run2.font.color.rgb = DARK

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("คู่มือการใช้งานโปรแกรม")
    run3.font.size = Pt(16)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("\n\nเว็บไซต์ใช้งาน: https://tdss-gules.vercel.app")
    run4.font.size = Pt(11)

    doc.add_page_break()


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT if level == 1 else DARK


def add_body(text):
    doc.add_paragraph(text)


def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_role_table():
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 2"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    headers = ["สิทธิ์", "ดูข้อมูล", "สร้าง/แก้ไข/อนุมัติงาน", "จัดการผู้ใช้งานองค์กร", "ตั้งค่าองค์กร"]
    for i, h in enumerate(headers):
        hdr[i].text = h
    rows = [
        ["Viewer", "✅", "❌", "❌", "❌"],
        ["Planner", "✅", "✅", "❌", "❌"],
        ["Org Admin", "✅", "✅", "✅", "✅"],
        ["System Owner", "✅ (ภาพรวมทุกองค์กร)", "❌", "✅ (ทั้งระบบ)", "✅ (ระดับระบบ)"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    doc.add_paragraph()


# ---------------------------------------------------------------------------
add_title_page()

add_heading("1. เกี่ยวกับ TDSS", level=1)
add_body(
    "TDSS (ระบบวางแผนงานขนส่ง) เป็นระบบสนับสนุนการตัดสินใจสำหรับวางแผนงานขนส่งโลจิสติกส์ "
    "ช่วยให้องค์กรขนส่งสร้างงานขนส่ง วางแผนเส้นทางและยานพาหนะ และรับคำแนะนำทางเลือกที่เหมาะสมที่สุด "
    "โดยใช้หลักการ AHP (Analytic Hierarchy Process) ซึ่งเป็นวิธีการเปรียบเทียบคู่ปัจจัยต่าง ๆ "
    "(เช่น ต้นทุน เวลา การปล่อยคาร์บอน) เพื่อคำนวณน้ำหนักความสำคัญและจัดอันดับทางเลือกอย่างเป็นระบบ"
)
add_body("ระบบรองรับการใช้งานแบบหลายองค์กร (Multi-tenant) โดยข้อมูลของแต่ละองค์กรแยกจากกันโดยสมบูรณ์")

add_heading("2. ระดับสิทธิ์ผู้ใช้งาน", level=1)
add_body("ระบบมี 4 ระดับสิทธิ์ ดังนี้")
add_role_table()
add_bullets([
    "Viewer — ดูข้อมูลอย่างเดียว เหมาะสำหรับผู้บริหารหรือผู้ที่ต้องการติดตามรายงาน",
    "Planner — พนักงานวางแผนตัวจริง สร้างงาน วางแผนเส้นทาง ใช้ AHP แนะนำ และอนุมัติแผนได้",
    "Org Admin — ดูแลองค์กร ทำได้ทุกอย่างของ Planner บวกจัดการสมาชิกและตั้งค่าองค์กร",
    "System Owner — ผู้ดูแลแพลตฟอร์มทั้งระบบ ดูแลทุกองค์กรจากภายนอก ไม่ยุ่งกับงานขนส่งในองค์กรโดยตรง",
])
doc.add_page_break()

add_heading("3. การเข้าสู่ระบบและสมัครใช้งาน", level=1)
add_body(
    "เปิดเว็บไซต์ https://tdss-gules.vercel.app จะพบหน้าเข้าสู่ระบบ กรอกอีเมลและรหัสผ่านที่ได้รับ "
    "จากผู้ดูแลองค์กร แล้วกดปุ่ม \"เข้าสู่ระบบ\""
)
add_steps([
    "เปิดเว็บไซต์ https://tdss-gules.vercel.app",
    "กรอกอีเมลในช่อง \"อีเมล\" และรหัสผ่านในช่อง \"รหัสผ่าน\"",
    "กดปุ่ม \"เข้าสู่ระบบ\" สีแดง",
    "ระบบจะพาไปยังแดชบอร์ดตามสิทธิ์ของบัญชีนั้น (องค์กร หรือ System Owner)",
])
add_body(
    "หากยังไม่มีองค์กร สามารถกดปุ่ม \"สมัครใช้งานองค์กรใหม่\" เพื่อสร้างองค์กรใหม่และบัญชี "
    "Org Admin ของตนเองได้ทันที โดยกรอกชื่อองค์กร ชื่อผู้ใช้ อีเมล และตั้งรหัสผ่าน"
)

doc.add_page_break()
add_heading("4. คู่มือฝั่งองค์กร (Org Workspace)", level=1)
add_body("เมนูในส่วนนี้ใช้งานโดยผู้ใช้ในองค์กร (Org Admin / Planner / Viewer)")

add_heading("4.1 แดชบอร์ด", level=2)
add_body("แสดงภาพรวมสถิติขององค์กร เช่น จำนวนงานขนส่ง สถานะงาน และคำแนะนำที่สร้างล่าสุด เลือกช่วงวันที่ดูข้อมูลย้อนหลังได้ที่มุมขวาบนของหน้า")

add_heading("4.2 งานขนส่ง (Jobs)", level=2)
add_body("แสดงรายการงานขนส่งทั้งหมดขององค์กร พร้อมสถานะ (รอวางแผน / วางแผนแล้ว / อนุมัติแล้ว ฯลฯ)")
add_steps([
    "กด \"+ สร้างงานขนส่ง\" เพื่อสร้างงานใหม่ กรอกต้นทาง ปลายทาง สินค้า และกำหนดส่ง",
    "กดเข้าไปที่แถวงานเพื่อดูรายละเอียด หรือกด \"วางแผน\" เพื่อเริ่ม Planning Wizard",
])

add_heading("4.3 ยานพาหนะ", level=2)
add_body("จัดการข้อมูลยานพาหนะขององค์กร (ทะเบียน ประเภท ความจุ) เพิ่ม/แก้ไข/ลบได้ (สิทธิ์ Planner ขึ้นไป)")

add_heading("4.4 เส้นทาง", level=2)
add_body("จัดการข้อมูลเส้นทางขนส่ง เลือกโหมดกำหนดเส้นทางเองหรืออิงจาก Google Maps ตามค่าเริ่มต้นขององค์กร")

add_heading("4.5 โปรไฟล์การตัดสินใจ (AHP)", level=2)
add_body(
    "จุดสำคัญของระบบ — ใช้กำหนดน้ำหนักความสำคัญของแต่ละปัจจัย (ต้นทุน, เวลา, CO2 ฯลฯ) ผ่านการเปรียบเทียบคู่ทีละคู่ "
    "ระบบจะคำนวณน้ำหนักและตรวจสอบความสมเหตุสมผล (Consistency Ratio) ให้อัตโนมัติ หากค่าความสอดคล้องไม่ผ่านเกณฑ์ "
    "โปรไฟล์จะถูกบันทึกเป็นฉบับร่างจนกว่าจะแก้ไขให้ผ่านเกณฑ์ สามารถแก้ไขหรือลบโปรไฟล์ได้ "
    "(ระบบป้องกันการลบโปรไฟล์ที่มีการใช้งานในแผนงานที่มีอยู่แล้ว)"
)
add_steps([
    "กดปุ่ม \"+ สร้างโปรไฟล์ใหม่\" แล้วตั้งชื่อโปรไฟล์และเลือกปัจจัยที่จะใช้เปรียบเทียบ",
    "เปรียบเทียบความสำคัญของปัจจัยทีละคู่ (เช่น \"ต้นทุนสำคัญกว่าเวลาแค่ไหน\") ตามสเกล 1-9",
    "ระบบคำนวณน้ำหนักของแต่ละปัจจัยและค่า Consistency Ratio (CR) ให้อัตโนมัติ",
    "ถ้า CR ผ่านเกณฑ์ (≤ 0.10) โปรไฟล์จะพร้อมใช้งานทันที ถ้าไม่ผ่านจะถูกบันทึกเป็นฉบับร่าง ให้กลับไปแก้ค่าเปรียบเทียบใหม่",
    "แก้ไขหรือลบโปรไฟล์ที่มีอยู่ได้จากปุ่มในตาราง (ลบไม่ได้หากมีแผนงานอ้างอิงโปรไฟล์นั้นอยู่)",
])

add_heading("4.6 สร้างแผนงาน (Planning Wizard)", level=2)
add_body("ขั้นตอนการวางแผนงานขนส่งแบบ Wizard 7 ขั้นตอน เข้าถึงได้จากหน้า \"งานขนส่ง\" โดยกดเลือกงานแล้วกด \"วางแผน\"")
add_steps([
    "ขั้นที่ 1: ยืนยันรายละเอียดงานขนส่ง (ต้นทาง ปลายทาง สินค้า กำหนดส่ง)",
    "ขั้นที่ 2: เลือกยานพาหนะที่เป็นไปได้ (ระบบกรองเฉพาะคันที่ว่างและเหมาะสม)",
    "ขั้นที่ 3: เลือกเส้นทางที่เป็นไปได้",
    "ขั้นที่ 4: เลือกโปรไฟล์ AHP ที่จะใช้ให้คะแนน (ระบบเลือกโปรไฟล์เริ่มต้นขององค์กรไว้ให้ล่วงหน้า)",
    "ขั้นที่ 5-6: ตรวจสอบเงื่อนไข/ข้อจำกัดเพิ่มเติม",
    "ขั้นที่ 7: ตรวจทานสรุปแผนทั้งหมดแล้วกด \"ยืนยันและสร้างคำแนะนำ\"",
])

add_heading("4.7 ผลลัพธ์คำแนะนำ", level=2)
add_body(
    "แสดงอันดับทางเลือก (คู่ยานพาหนะ-เส้นทาง) พร้อมคะแนนรวมและคำอธิบายเหตุผลแบบภาษาที่เข้าใจง่าย"
)
add_steps([
    "ดูตารางอันดับทางเลือกเรียงจากคะแนนสูงสุดไปต่ำสุด พร้อมคำอธิบายว่าทำไมทางเลือกนั้นได้คะแนนดี/ไม่ดี",
    "กดเลือกทางเลือกที่ต้องการ",
    "กดปุ่ม \"อนุมัติแผนนี้\" เพื่อยืนยัน สถานะงานจะเปลี่ยนเป็น \"อนุมัติแล้ว\" และมีการแจ้งเตือนไปยังสมาชิกในองค์กร (ถ้าเปิดใช้งานไว้)",
])

add_heading("4.8 รายงาน", level=2)
add_body(
    "รายงานสรุปด้านต่าง ๆ เช่น การใช้ยานพาหนะ เปรียบเทียบต้นทุน ปริมาณ CO2 และการใช้งานโปรไฟล์ AHP"
)
add_steps([
    "เลือกช่วงวันที่ที่ต้องการดูข้อมูลจากตัวเลือก 7 วัน / 30 วัน / 90 วัน / ทั้งหมด (หรือกำหนดเอง)",
    "เลือกแท็บรายงานที่ต้องการดู (การใช้ยานพาหนะ / ต้นทุน / CO2 / โปรไฟล์ AHP)",
    "กดปุ่ม \"ดาวน์โหลด CSV\" เพื่อบันทึกรายงานเป็นไฟล์ (การดาวน์โหลดทุกครั้งจะถูกบันทึกใน audit log อัตโนมัติ)",
])

add_heading("4.9 ผู้ใช้งานองค์กร (Org Admin เท่านั้น)", level=2)
add_body("จัดการสมาชิกในองค์กร")
add_steps([
    "ดูตารางสมาชิกทั้งหมดพร้อมบทบาทปัจจุบัน",
    "กด \"+ เพิ่มผู้ใช้งาน\" กรอกชื่อ อีเมล รหัสผ่าน และเลือกบทบาท (org_admin / planner / viewer) เพื่อเพิ่มสมาชิกใหม่",
    "เปลี่ยนบทบาทสมาชิกที่มีอยู่ได้ทันทีจาก dropdown ในตาราง",
])

add_heading("4.10 ข้อมูลองค์กร", level=2)
add_body("แสดง/แก้ไขข้อมูลพื้นฐานขององค์กร เช่น ชื่อ รหัสองค์กร ที่อยู่ ช่องทางติดต่อ ผู้ใช้ทุกบทบาทดูได้ แต่แก้ไขได้เฉพาะ Org Admin")

add_heading("4.11 ตั้งค่าองค์กร (Org Admin เท่านั้น)", level=2)
add_body(
    "กำหนดค่าเริ่มต้นขององค์กร ได้แก่ โปรไฟล์ AHP เริ่มต้นที่จะถูกเลือกไว้ล่วงหน้าในแผนงานใหม่ "
    "โหมดเส้นทางเริ่มต้น (กำหนดเอง / อิง Google Maps) และสวิตช์เปิด/ปิดการแจ้งเตือนอัตโนมัติ "
    "(เมื่อคำแนะนำเสร็จ / เมื่องานถูกอนุมัติ) แก้ไขแล้วกด \"บันทึก\" เพื่อให้มีผลทันที"
)

add_heading("4.12 โปรไฟล์ของฉัน", level=2)
add_body("แก้ไขชื่อผู้ใช้และเปลี่ยนรหัสผ่านของบัญชีตนเองได้จากหน้านี้")

doc.add_page_break()
add_heading("5. คู่มือฝั่ง System Owner Console", level=1)
add_body("ส่วนนี้ใช้งานโดยบัญชี System Owner เท่านั้น สำหรับดูแลแพลตฟอร์มทั้งหมดในภาพรวม")

add_heading("5.1 แดชบอร์ดระบบ", level=2)
add_body("ภาพรวมทั้งแพลตฟอร์ม จำนวนองค์กรทั้งหมด องค์กรที่ใช้งานอยู่/ถูกระงับ ผู้ใช้งานทั้งหมด งานขนส่งทั้งหมด และคำแนะนำที่สร้างแล้ว เป็นหน้าแรกที่พบทันทีหลัง System Owner ล็อกอิน")

add_heading("5.2 องค์กรทั้งหมด", level=2)
add_body("รายการองค์กรทั้งหมดในระบบ กดเข้าไปดูรายละเอียดแต่ละองค์กร และกดปุ่มระงับ/เปิดใช้งานองค์กรได้ (เมื่อระงับ สมาชิกในองค์กรนั้นจะเข้าระบบไม่ได้ชั่วคราว)")

add_heading("5.3 ผู้ใช้งานทั้งหมด", level=2)
add_body("รายชื่อผู้ใช้งานทุกคนในทุกองค์กร พร้อมองค์กรที่สังกัดและบทบาท กดปุ่ม \"ระงับ\" ท้ายแถวเพื่อปิดการใช้งานบัญชีรายบุคคลได้ทันที")

add_heading("5.4 การใช้งาน", level=2)
add_body(
    "สถิติการใช้งานจริงของแต่ละองค์กร ได้แก่ จำนวนงานขนส่ง จำนวนคำแนะนำที่สร้าง ผู้ใช้งานที่ยัง active "
    "จำนวนยานพาหนะ เส้นทาง และจำนวนรายงานที่ถูก export กรองตามชื่อองค์กรหรือช่วงวันที่ได้"
)

add_heading("5.5 บันทึกการใช้งาน (Audit Log)", level=2)
add_body("บันทึกประวัติการกระทำสำคัญทั้งหมดในระบบ (สร้าง/แก้ไข/ลบ/อนุมัติ/ดาวน์โหลดรายงาน) พร้อมผู้กระทำและเวลา ใช้สำหรับตรวจสอบย้อนหลัง")

add_heading("5.6 สถานะระบบ", level=2)
add_body("แสดงสถานะสุขภาพของระบบโดยรวม สำหรับให้ System Owner ตรวจสอบว่าแพลตฟอร์มทำงานปกติ")

add_heading("5.7 ตั้งค่าระบบ", level=2)
add_body("กำหนดค่าระดับแพลตฟอร์มที่มีผลต่อผู้ใช้งานทุกองค์กร ได้แก่ ชื่อแอปที่แสดงผล และข้อความแบนเนอร์แจ้งเตือน (เช่น ประกาศปิดปรับปรุงระบบ) แก้ไขแล้วกด \"บันทึก\" มีผลทันทีทุกองค์กร")

add_heading("5.8 โปรไฟล์ของฉัน", level=2)
add_body("แก้ไขชื่อและเปลี่ยนรหัสผ่านของบัญชี System Owner")

doc.add_page_break()
add_heading("ภาคผนวก: บัญชีทดสอบ", level=1)
add_body("สำหรับทดลองใช้งานระบบ (รหัสผ่านทุกบัญชี: password123)")
add_bullets([
    "System Owner: owner@tdss.local",
    "Org 1 — Admin: admin@orgone.local / Planner: planner@orgone.local / Viewer: viewer@orgone.local",
    "Org 2 — Admin: admin@orgtwo.local / Planner: planner@orgtwo.local / Viewer: viewer@orgtwo.local",
])

os.makedirs(BASE_DIR, exist_ok=True)
doc.save(OUT_PATH)
print("Saved:", OUT_PATH)
